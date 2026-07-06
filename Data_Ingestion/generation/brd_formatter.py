"""
BRD Document Formatter
=======================
Produces a styled Business Requirements Document (.docx) from AI-generated
section content, matching the Adani BRD template (Blockchain Platform.docx).

Structure: 13 numbered top-level sections, 23 generated content blocks,
6 table-based requirement sections.

Entry point:  format_brd_docx(sections_by_key, project_name, out_path)
Called from:  doc_writer.export_job() when doc_type is BRD and format is DOCX.
"""
from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Optional

# ─────────────────────────────────────────────────────────────────────────────
# Adani brand colours (hex strings, no #)
# ─────────────────────────────────────────────────────────────────────────────
_DARK_BLUE   = "1F3763"   # H3 text, table header fill
_MID_BLUE    = "2E75B6"   # header rule line
_LIGHT_BLUE  = "D6E4F7"   # table alt-row fill (unused by default, kept for future)
_WHITE       = "FFFFFF"
_BLACK       = "000000"
_GRAY        = "595959"   # sub-text, footer

# ─────────────────────────────────────────────────────────────────────────────
# Document structure
# ─────────────────────────────────────────────────────────────────────────────
# key=None entries are pure heading rows — the formatter injects them as headings
# with no AI content.  key=<str> entries pull from sections_by_key dict.
BRD_STRUCTURE = [
    {"num": "1",      "key": "purpose",                      "title": "Purpose",                                                        "level": 1},
    {"num": "2",      "key": "scope",                        "title": "Scope",                                                          "level": 1},
    {"num": "3",      "key": "terms_and_definitions",        "title": "Terms and Definition",                                           "level": 1},
    {"num": "4",      "key": None,                           "title": "Business requirements summary",                                  "level": 1},
    {"num": "4.1",    "key": "business_need",                "title": "Business need and why?",                                         "level": 2},
    {"num": "4.2",    "key": "business_context",             "title": "Context for business need",                                      "level": 2},
    {"num": "4.3",    "key": "problem_statement",            "title": "Problem statement",                                              "level": 2},
    {"num": "4.4",    "key": "business_value",               "title": "Business value delivered",                                       "level": 2},
    {"num": "4.5",    "key": None,                           "title": "Competitor Landscape and Positioning",                           "level": 2},
    {"num": "4.5.1",  "key": "competitor_landscape",         "title": "Competitor landscape",                                           "level": 3},
    {"num": "4.5.2",  "key": "similar_digital_solutions",    "title": "Similar digital solutions",                                      "level": 3},
    {"num": "4.5.3",  "key": "positioning",                  "title": "Positioning",                                                    "level": 3},
    {"num": "4.6",    "key": "critical_success_factors",     "title": "Critical success factors",                                       "level": 2},
    {"num": "5",      "key": "business_process_overview",    "title": "Business process",                                               "level": 1},
    {"num": "5.1",    "key": "pain_areas_operational",       "title": "Pain areas or opportunities in existing operational process",    "level": 2},
    {"num": "5.2",    "key": "pain_areas_technology",        "title": "Pain areas or opportunities in existing technology/systems",     "level": 2},
    {"num": "5.3",    "key": "business_functionality_impact","title": "Business functionality impact",                                  "level": 2},
    {"num": "5.4",    "key": None,                           "title": "Business process impacted",                                      "level": 2},
    {"num": "5.4.1",  "key": "as_is_process",                "title": "As-Is Business Process",                                        "level": 3},
    {"num": "5.4.2",  "key": "to_be_process",                "title": "To-Be Business Process (New/changes only)",                     "level": 3},
    {"num": "6",      "key": "business_use_cases",           "title": "Business use cases",                                            "level": 1},
    {"num": "7",      "key": "business_requirements_list",   "title": "Business requirements",                                         "level": 1},
    {"num": "8",      "key": "functional_requirements",      "title": "Functional requirements",                                        "level": 1},
    {"num": "9",      "key": "non_functional_requirements",  "title": "Non-functional requirements",                                    "level": 1},
    {"num": "10",     "key": "solution_overview",            "title": "Solution overview-Cloud first approach",                        "level": 1},
    {"num": "11",     "key": "key_constraints",              "title": "Key constraints",                                                "level": 1},
    {"num": "12",     "key": "project_schedule",             "title": "Project schedule",                                               "level": 1},
    {"num": "13",     "key": "other_constraints_assumptions","title": "Other constraints and assumptions if any",                       "level": 1},
]


# ─────────────────────────────────────────────────────────────────────────────
# Public entry point
# ─────────────────────────────────────────────────────────────────────────────

def format_brd_docx(
    sections_by_key: dict[str, str],
    project_name: str,
    out_path: Path,
    *,
    client_name: str = "Adani Energy Solutions",
    doc_version: str = "1.0",
) -> None:
    """
    Build a styled BRD Word document from AI-generated section content.

    Args:
        sections_by_key:  {section_key: markdown_content_string}
        project_name:     Used in header, cover, and document control table.
        out_path:         Destination .docx path (must include .docx extension).
        client_name:      Override the client name shown on cover / control table.
        doc_version:      Document version string (e.g. "1.0").
    """
    from docx import Document

    doc = Document()
    _setup_page(doc)
    _setup_styles(doc)
    _add_header(doc, project_name)
    _add_footer(doc)
    _add_cover(doc, project_name, client_name, doc_version)
    _add_all_sections(doc, sections_by_key)
    doc.save(str(out_path))


# ─────────────────────────────────────────────────────────────────────────────
# Generic structured formatter — used by every NON-BRD document type
# (RFP, SOW, NIT, BOQ, NFA, ARB, Proposal, Tech Spec, Scope, …)
# Reuses the exact BRD styling (page, fonts, header+logo, footer, styled
# headings, dark-blue table headers) but is driven by the document's own
# ordered section list rather than the fixed BRD numbering.
# ─────────────────────────────────────────────────────────────────────────────

def format_structured_docx(
    doc_type: str,
    project_name: str,
    sections: list[dict],
    out_path: Path,
    *,
    client_name: str = "Adani Energy Solutions",
    doc_version: str = "1.0",
) -> None:
    """
    Build a styled Word document for any document type from AI-generated content.

    Args:
        doc_type:      Display name (e.g. "Request for Proposal (RFP)").
        project_name:  Shown on the cover, header, and document-control table.
        sections:      Ordered list of {"title": str, "content": markdown str}.
        out_path:      Destination .docx path (must include .docx).
    """
    from docx import Document

    doc = Document()
    _setup_page(doc)
    _setup_styles(doc)
    _add_header(doc, project_name, label=_short_doc_label(doc_type))
    _add_footer(doc)
    _add_generic_cover(doc, doc_type, project_name, client_name, doc_version)
    _add_generic_sections(doc, sections)
    doc.save(str(out_path))


def _short_doc_label(doc_type: str) -> str:
    """Derive a short header label — the parenthesised abbreviation if present
    (e.g. 'Request for Proposal (RFP)' → 'RFP'), else the doc type itself."""
    m = re.search(r"\(([A-Za-z]{2,6})\)", doc_type or "")
    if m:
        return m.group(1).upper()
    return (doc_type or "Document").strip()


def _add_generic_cover(doc, doc_type: str, project_name: str,
                       client_name: str, doc_version: str) -> None:
    """Lightweight cover: logo, title (= doc type), subtitle (= project),
    a document-control table, and a confidentiality note."""
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    today = datetime.utcnow().strftime("%d-%m-%Y")
    year  = datetime.utcnow().strftime("%Y")
    abbr  = _short_doc_label(doc_type)
    safe_id = re.sub(r"[^\w]", "_", project_name)[:30]
    document_id = f"{year}_AESL_{safe_id}_{abbr}_001"

    logo_path = Path(__file__).parent.parent / "static" / "adani_logo.jpg"
    if logo_path.exists():
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_before = Pt(24)
        lp.paragraph_format.space_after  = Pt(8)
        lp.add_run().add_picture(str(logo_path), width=Cm(5))

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(12)
    t.paragraph_format.space_after  = Pt(4)
    r = t.add_run(doc_type)
    r.font.name = "Calibri"
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x37, 0x63)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(4)
    r2 = sub.add_run(project_name)
    r2.font.name = "Calibri"
    r2.font.size = Pt(14)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph()

    _add_cover_section_heading(doc, "Document Control")
    ctrl = doc.add_table(rows=4, cols=2)
    ctrl.style = _safe_table_style(doc)
    pairs = [
        ("Document ID:", document_id),
        ("Client:",      client_name),
        ("Version:",     doc_version),
        ("Date:",        today),
    ]
    for i, (k, v) in enumerate(pairs):
        _set_cell_text(ctrl.rows[i].cells[0], k, bold=True, bg=_DARK_BLUE, fg=_WHITE)
        _set_cell_text(ctrl.rows[i].cells[1], v)
    _set_col_widths(ctrl, [Cm(4.5), Cm(11.42)])

    doc.add_paragraph()

    _add_cover_section_heading(doc, "Confidentiality")
    conf = doc.add_paragraph()
    cr = conf.add_run(
        "This document contains restricted information pertaining to Adani. "
        "The addressee should honour these access rights by preventing intentional "
        "or accidental access outside the intended scope."
    )
    cr.font.name = "Calibri"
    cr.font.size = Pt(10)
    cr.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_page_break()


def _add_generic_sections(doc, sections: list[dict]) -> None:
    """Render each section as a numbered H1 heading + styled markdown body."""
    n = 0
    for sec in sections:
        content = (sec.get("content") or "").strip()
        if not content:
            continue   # skip empty sections so the DOCX matches the preview
        n += 1
        title = sec.get("title") or f"Section {n}"
        doc.add_heading(f"{n}. {title}", level=1)
        _render_brd_markdown(doc, content)
        doc.add_paragraph()   # spacer between sections


# ─────────────────────────────────────────────────────────────────────────────
# Page and style setup
# ─────────────────────────────────────────────────────────────────────────────

def _setup_page(doc) -> None:
    from docx.shared import Cm
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.54)
    sec.right_margin  = Cm(2.54)
    sec.top_margin    = Cm(2.54)
    sec.bottom_margin = Cm(2.0)


def _setup_styles(doc) -> None:
    from docx.shared import Pt, RGBColor

    styles = doc.styles

    # Normal body text
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Heading 1: 16pt, black, bold
    h1 = styles["Heading 1"]
    h1.font.name  = "Calibri"
    h1.font.size  = Pt(16)
    h1.font.bold  = True
    h1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h1.font.underline = False
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after  = Pt(6)
    h1.paragraph_format.keep_with_next = True

    # Heading 2: 13pt, black, bold
    h2 = styles["Heading 2"]
    h2.font.name  = "Calibri"
    h2.font.size  = Pt(13)
    h2.font.bold  = True
    h2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h2.font.underline = False
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after  = Pt(4)
    h2.paragraph_format.keep_with_next = True

    # Heading 3: 12pt, Adani dark blue, bold
    h3 = styles["Heading 3"]
    h3.font.name  = "Calibri"
    h3.font.size  = Pt(12)
    h3.font.bold  = True
    h3.font.color.rgb = RGBColor(0x1F, 0x37, 0x63)
    h3.font.underline = False
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after  = Pt(3)
    h3.paragraph_format.keep_with_next = True


# ─────────────────────────────────────────────────────────────────────────────
# Header and footer
# ─────────────────────────────────────────────────────────────────────────────

def _add_header(doc, project_name: str, label: str = "Detailed BRD") -> None:
    """Header matches the Drishti BRD reference: two separate right-aligned
    paragraphs so the logo is guaranteed to land at the top-right corner.
    Tab-stop approaches are unreliable for inline images in Word headers.

    `label` prefixes the header text (e.g. "Detailed BRD" for BRD, or the
    document abbreviation like "RFP" for other document types)."""
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement

    section = doc.sections[0]
    header  = section.header

    # ── Paragraph 1: right-aligned BRD title with blue bottom border ──────────
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    # Clear any leftover text/runs from a prior call
    for run in hp.runs:
        run.text = ""

    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after  = Pt(2)
    _set_para_border_bottom(hp, _MID_BLUE, size=6)

    run_txt = hp.add_run(f"{label}  |  {project_name}")
    run_txt.font.name      = "Calibri"
    run_txt.font.size      = Pt(9)
    run_txt.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    run_txt.font.italic    = True

    # ── Paragraph 2: right-aligned Adani logo (separate paragraph = top-right) ─
    logo_path = Path(__file__).parent.parent / "static" / "adani_logo.jpg"
    if logo_path.exists():
        hp2 = header.add_paragraph()
        hp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp2.paragraph_format.space_before = Pt(0)
        hp2.paragraph_format.space_after  = Pt(0)

        run_logo = hp2.add_run()
        inline   = run_logo.add_picture(str(logo_path), height=Cm(0.9))
        ns_a = "http://schemas.openxmlformats.org/drawingml/2006/main"
        blip = inline._inline.find(f".//{{{ns_a}}}blip")
        if blip is not None:
            alpha_fix = OxmlElement("a:alphaModFix")
            alpha_fix.set("amt", "90000")  # 90 000 per-mille = 90% opacity
            blip.append(alpha_fix)


def _add_footer(doc) -> None:
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    section = doc.sections[0]
    footer  = section.footer

    for para in footer.paragraphs:
        for run in para.runs:
            run.text = ""
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # "Page X of Y" using Word field codes
    run = fp.add_run()
    run.font.name  = "Calibri"
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    _append_field(run, "PAGE")

    run2 = fp.add_run(" of ")
    run2.font.name  = "Calibri"
    run2.font.size  = Pt(9)
    run2.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    run3 = fp.add_run()
    run3.font.name  = "Calibri"
    run3.font.size  = Pt(9)
    run3.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    _append_field(run3, "NUMPAGES")


# ─────────────────────────────────────────────────────────────────────────────
# Cover page  (title + document control table + revision history)
# ─────────────────────────────────────────────────────────────────────────────

def _add_cover(doc, project_name: str, client_name: str, doc_version: str) -> None:
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    today = datetime.utcnow().strftime("%d-%m-%Y")

    # ── Document ID derived from project name ─────────────────────────────────
    year       = datetime.utcnow().strftime("%Y")
    safe_id    = re.sub(r"[^\w]", "_", project_name)[:30]
    document_id = f"{year}_AESL_{safe_id}_001"

    # ── Adani logo ────────────────────────────────────────────────────────────
    logo_path = Path(__file__).parent.parent / "static" / "adani_logo.jpg"
    if logo_path.exists():
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.paragraph_format.space_before = Pt(24)
        logo_p.paragraph_format.space_after  = Pt(8)
        logo_p.add_run().add_picture(str(logo_path), width=Cm(5))

    # ── Main title ────────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(12)
    t.paragraph_format.space_after  = Pt(4)
    run = t.add_run("Detailed BRD for AESL")
    run.font.name  = "Calibri"
    run.font.size  = Pt(20)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x1F, 0x37, 0x63)

    # Sub-title (project name)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(4)
    r2 = sub.add_run(project_name)
    r2.font.name  = "Calibri"
    r2.font.size  = Pt(14)
    r2.font.bold  = True
    r2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph()  # spacer

    # ── Document Control block (Document ID + Version) ────────────────────────
    _add_cover_section_heading(doc, "Document Control")
    ctrl_tbl = doc.add_table(rows=2, cols=2)
    ctrl_tbl.style = _safe_table_style(doc)
    _set_cell_text(ctrl_tbl.rows[0].cells[0], "Document ID:", bold=True, bg=_DARK_BLUE, fg=_WHITE)
    _set_cell_text(ctrl_tbl.rows[0].cells[1], document_id)
    _set_cell_text(ctrl_tbl.rows[1].cells[0], "Version:", bold=True, bg=_DARK_BLUE, fg=_WHITE)
    _set_cell_text(ctrl_tbl.rows[1].cells[1], doc_version)
    _set_col_widths(ctrl_tbl, [Cm(4.5), Cm(11.42)])

    doc.add_paragraph()

    # ── Author table ─────────────────────────────────────────────────────────
    _add_cover_section_heading(doc, "Document Author")
    auth_tbl = doc.add_table(rows=2, cols=4)
    auth_tbl.style = _safe_table_style(doc)
    for c_idx, hdr in enumerate(["#", "Name of Author", "Role", "Signature"]):
        _set_cell_text(auth_tbl.rows[0].cells[c_idx], hdr, bold=True, bg=_DARK_BLUE, fg=_WHITE)
    for c_idx, val in enumerate(["1", "Intellidraft AI", "Document Author", ""]):
        _set_cell_text(auth_tbl.rows[1].cells[c_idx], val)
    _set_col_widths(auth_tbl, [Cm(1), Cm(5), Cm(5), Cm(4.92)])

    doc.add_paragraph()

    # ── Reviewer / Approver table ─────────────────────────────────────────────
    _add_cover_section_heading(doc, "Reviewer / Approver")
    rev_tbl2 = doc.add_table(rows=2, cols=5)
    rev_tbl2.style = _safe_table_style(doc)
    for c_idx, hdr in enumerate(["#", "Reviewer / Approver", "Name", "Role", "Signature"]):
        _set_cell_text(rev_tbl2.rows[0].cells[c_idx], hdr, bold=True, bg=_DARK_BLUE, fg=_WHITE)
    for c_idx, val in enumerate(["1", "Reviewer", "[To be assigned]", "[To be assigned]", ""]):
        _set_cell_text(rev_tbl2.rows[1].cells[c_idx], val)
    _set_col_widths(rev_tbl2, [Cm(1), Cm(3), Cm(4), Cm(4), Cm(3.92)])

    doc.add_paragraph()

    # ── Revision History ──────────────────────────────────────────────────────
    _add_cover_section_heading(doc, "Revision History")
    rev_tbl = doc.add_table(rows=2, cols=3)
    rev_tbl.style = _safe_table_style(doc)
    for c_idx, hdr in enumerate(["Revision", "Date of Change", "Revision Description"]):
        _set_cell_text(rev_tbl.rows[0].cells[c_idx], hdr, bold=True, bg=_DARK_BLUE, fg=_WHITE)
    for c_idx, val in enumerate([doc_version, today, "Intellidraft AI, Initial Draft version"]):
        _set_cell_text(rev_tbl.rows[1].cells[c_idx], val)
    _set_col_widths(rev_tbl, [Cm(2), Cm(3), Cm(10.92)])

    doc.add_paragraph()

    # ── Referenced Documents ──────────────────────────────────────────────────
    _add_cover_section_heading(doc, "Referenced Documents")
    ref_tbl = doc.add_table(rows=2, cols=3)
    ref_tbl.style = _safe_table_style(doc)
    for c_idx, hdr in enumerate(["Ref No.", "Document ID", "Document Title"]):
        _set_cell_text(ref_tbl.rows[0].cells[c_idx], hdr, bold=True, bg=_DARK_BLUE, fg=_WHITE)
    for c_idx, val in enumerate(["-", "-", "-"]):
        _set_cell_text(ref_tbl.rows[1].cells[c_idx], val)
    _set_col_widths(ref_tbl, [Cm(2), Cm(4), Cm(9.92)])

    doc.add_paragraph()

    # ── Confidentiality ───────────────────────────────────────────────────────
    _add_cover_section_heading(doc, "Confidentiality")
    conf = doc.add_paragraph()
    cr = conf.add_run(
        "This document contains restricted information pertaining to Adani. "
        "The access level for the document is specified above. "
        "The addressee should honour this access rights by preventing intentional or "
        "accidental access outside the access scope."
    )
    cr.font.name  = "Calibri"
    cr.font.size  = Pt(10)
    cr.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    # ── Disclaimer ────────────────────────────────────────────────────────────
    _add_cover_section_heading(doc, "Disclaimer")
    disc = doc.add_paragraph()
    dr = disc.add_run(
        "This document is solely for the information of Adani and should not be used, "
        "circulated, quoted or otherwise referred to for any other purpose, nor included or "
        "referred to in whole or in part in any document without our prior written consent."
    )
    dr.font.name  = "Calibri"
    dr.font.size  = Pt(10)
    dr.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Page break to start content on a new page
    doc.add_page_break()


def _add_cover_section_heading(doc, text: str) -> None:
    """Small bold heading used inside the cover page sections."""
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(11)
    r.font.bold  = True
    r.font.color.rgb = RGBColor(0x1F, 0x37, 0x63)


def _set_col_widths(tbl, widths: list) -> None:
    """Set column widths on every cell in a table (workaround for python-docx width precedence)."""
    for row in tbl.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = w


# ─────────────────────────────────────────────────────────────────────────────
# Section rendering
# ─────────────────────────────────────────────────────────────────────────────

def _add_all_sections(doc, sections_by_key: dict[str, str]) -> None:
    for entry in BRD_STRUCTURE:
        num   = entry["num"]
        key   = entry["key"]
        title = entry["title"]
        level = entry["level"]

        sep = ". " if "." not in num else " "
        heading_text = f"{num}{sep}{title}"

        if key is None:
            # Structural parent heading — always include
            doc.add_heading(heading_text, level=level)
            continue

        content = sections_by_key.get(key, "").strip()
        if not content:
            # Skip heading AND content for un-generated sections so the DOCX
            # matches the preview exactly (both omit sections with no content)
            continue

        doc.add_heading(heading_text, level=level)
        _render_brd_markdown(doc, content)
        doc.add_paragraph()  # spacer after each section


# ─────────────────────────────────────────────────────────────────────────────
# Markdown renderer  (fork of doc_writer._render_markdown_to_docx with
# Adani-styled table headers)
# ─────────────────────────────────────────────────────────────────────────────

def _render_brd_markdown(doc, md_text: str) -> None:
    """
    Parse Markdown and add styled content to the Document.
    Handles: headings (## / ###), bold/italic/code spans, bullet lists,
    numbered lists, code blocks, horizontal rules, plain paragraphs,
    and pipe tables (with Adani-styled header row).
    """
    from docx.shared import Pt

    lines     = md_text.split("\n")
    i         = 0
    in_code   = False
    code_buf: list[str] = []
    table_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        # ── Code block ─────────────────────────────────────────────────
        if line.strip().startswith("```"):
            if in_code:
                if code_buf:
                    p   = doc.add_paragraph()
                    run = p.add_run("\n".join(code_buf))
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                code_buf = []
                in_code  = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ── Pipe table ──────────────────────────────────────────────────
        if line.strip().startswith("|"):
            table_buf.append(line)
            i += 1
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_buf.append(lines[i])
                i += 1
            _add_brd_table(doc, table_buf)
            table_buf = []
            continue

        # ── Headings (avoid re-adding headings the LLM may have prefixed) ──
        if line.startswith("### "):
            doc.add_heading(_strip_md(line[4:]), level=3)
        elif line.startswith("## "):
            doc.add_heading(_strip_md(line[3:]), level=2)
        elif line.startswith("# "):
            doc.add_heading(_strip_md(line[2:]), level=1)

        # ── Horizontal rule ────────────────────────────────────────────
        elif line.strip() in ("---", "***", "___"):
            pass  # skip HRs — section spacing handles visual separation

        # ── Bullet (- or *) ────────────────────────────────────────────
        elif line.startswith("- ") or line.startswith("* "):
            _inline_runs(doc.add_paragraph(style="List Bullet"), line[2:])

        # ── Numbered list ──────────────────────────────────────────────
        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line)
            _inline_runs(doc.add_paragraph(style="List Number"), text)

        # ── Sub-bullet (  - ) ──────────────────────────────────────────
        elif re.match(r"^\s{2,}[-*] ", line):
            _inline_runs(doc.add_paragraph(style="List Bullet 2"), line.strip()[2:])

        # ── **Bold heading** pattern used in mixed sections ────────────
        elif re.match(r"^\*\*.+\*\*$", line.strip()):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            _inline_runs(p, line.strip())

        # ── Empty line ─────────────────────────────────────────────────
        elif not line.strip():
            pass

        # ── Normal paragraph ───────────────────────────────────────────
        else:
            _inline_runs(doc.add_paragraph(), line)

        i += 1


# ─────────────────────────────────────────────────────────────────────────────
# Adani-styled table (dark blue header row)
# ─────────────────────────────────────────────────────────────────────────────

def _add_brd_table(doc, table_lines: list[str]) -> None:
    """Convert Markdown pipe table lines to a styled python-docx Table."""
    from docx.shared import Pt

    rows: list[list[str]] = []
    for raw in table_lines:
        if re.match(r"^\|[-| :]+\|$", raw.strip()):
            continue  # skip separator row
        cells = [c.strip() for c in raw.strip().strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    tbl      = doc.add_table(rows=len(rows), cols=max_cols)
    tbl.style = _safe_table_style(doc)

    for r_idx, row_cells in enumerate(rows):
        for c_idx in range(max_cols):
            cell_text = row_cells[c_idx] if c_idx < len(row_cells) else ""
            cell      = tbl.rows[r_idx].cells[c_idx]
            is_header = (r_idx == 0)
            _set_cell_text(
                cell,
                _strip_md(cell_text),
                bold=is_header,
                bg=_DARK_BLUE if is_header else None,
                fg=_WHITE    if is_header else None,
            )

    doc.add_paragraph()  # spacing after table


# ─────────────────────────────────────────────────────────────────────────────
# Cell helpers
# ─────────────────────────────────────────────────────────────────────────────

def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    bg: Optional[str] = None,
    fg: Optional[str] = None,
) -> None:
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    para = cell.paragraphs[0]
    if para.runs:
        run = para.runs[0]
        run.text = text
    else:
        run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.bold = bold
    if fg:
        r, g, b = int(fg[0:2], 16), int(fg[2:4], 16), int(fg[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)

    if bg:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  bg)
        tcPr.append(shd)


# ─────────────────────────────────────────────────────────────────────────────
# Inline Markdown runs (bold / italic / code)
# ─────────────────────────────────────────────────────────────────────────────

def _inline_runs(para, text: str) -> None:
    """Add runs to a paragraph, applying **bold**, *italic*, and `code` spans."""
    from docx.shared import Pt

    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    for part in pattern.split(text):
        if part.startswith("**") and part.endswith("**"):
            r      = para.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r        = para.add_run(part[1:-1])
            r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r           = para.add_run(part[1:-1])
            r.font.name = "Courier New"
            r.font.size = Pt(9)
        else:
            para.add_run(part)


def _strip_md(text: str) -> str:
    """Strip inline Markdown markers from a string."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*",     r"\1", text)
    text = re.sub(r"`(.+?)`",       r"\1", text)
    return text.strip()


# ─────────────────────────────────────────────────────────────────────────────
# XML / style helpers
# ─────────────────────────────────────────────────────────────────────────────

def _set_para_border_bottom(para, color_hex: str, size: int = 6) -> None:
    """Add a bottom border to a paragraph via OOXML pPr/pBdr."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(size))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)


def _append_field(run, field_name: str) -> None:
    """Append a simple Word field code (e.g. PAGE, NUMPAGES) to a run element."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.text = f" {field_name} "

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for el in (begin, instr, sep, end):
        run._r.append(el)


def _safe_table_style(doc):
    """Return 'Table Grid' if available, otherwise the first table style."""
    for name in ("Table Grid", "Normal Table", "Table Normal"):
        try:
            return doc.styles[name]
        except KeyError:
            continue
    return None
